#!/usr/bin/env python3
"""
Test script to verify emoji rendering in DOCX export.
"""

import emoji as emoji_lib
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_emoji_font(run):
    """Apply emoji font configuration"""
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
        rFonts.set(qn("w:eastAsia"), "Arial")
        rFonts.set(qn("w:cs"), "Arial")
        rFonts.set(qn("w:hint"), "eastAsia")

        existing_rFonts = rPr.find(qn("w:rFonts"))
        if existing_rFonts is not None:
            rPr.remove(existing_rFonts)
        rPr.insert(0, rFonts)
    except Exception as e:
        print(f"Error: {e}")


def add_text_with_emoji_support(paragraph, text: str):
    """Split text into emoji and non-emoji segments"""
    segments = []
    current_segment = ""
    is_emoji_segment = False

    for char in text:
        char_is_emoji = char in emoji_lib.EMOJI_DATA

        if char_is_emoji != is_emoji_segment:
            if current_segment:
                segments.append((is_emoji_segment, current_segment))
            current_segment = char
            is_emoji_segment = char_is_emoji
        else:
            current_segment += char

    if current_segment:
        segments.append((is_emoji_segment, current_segment))

    print(f"\nText: {text}")
    print(f"Segments: {segments}")

    for is_emoji, segment_text in segments:
        run = paragraph.add_run(segment_text)
        run.font.size = Pt(14)
        if is_emoji:
            set_emoji_font(run)
            print(f"  Emoji segment: {segment_text!r}")
        else:
            print(f"  Text segment: {segment_text!r}")


# Create test document
doc = Document()
doc.add_heading("Emoji Test Document", 0)

# Test cases
test_texts = [
    "Hello 👋 World",
    "修复完成 ✅",
    "🎉 Success!",
    "Test 🔧 with 💡 multiple 🚀 emoji",
    "纯文本没有emoji",
    "😀😁😂🤣😃😄😅😆😉😊",
]

for test_text in test_texts:
    p = doc.add_paragraph()
    add_text_with_emoji_support(p, test_text)

# Save document
output_path = "emoji_test.docx"
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")
print("Open this file in WPS Office/Microsoft Word to verify emoji display")
