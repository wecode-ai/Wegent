# SPDX-FileCopyrightText: 2025 Weibo, Inc.
#
# SPDX-License-Identifier: Apache-2.0

"""Tests for explicit Latin and East Asian fonts in DOCX exports."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.text.font import CT_RPr

from app.services.export.docx_generator import (
    DOCUMENT_LANGUAGE,
    EAST_ASIA_FONT,
    EAST_ASIA_LANGUAGE,
    LATIN_FONT,
    MONOSPACE_FONT,
    _add_inline_formatting,
    _setup_document_styles,
)


def _font_attributes(rpr: CT_RPr) -> dict[str, str]:
    r_fonts = rpr.find(qn("w:rFonts"))
    return {
        "ascii": r_fonts.get(qn("w:ascii")),
        "hAnsi": r_fonts.get(qn("w:hAnsi")),
        "eastAsia": r_fonts.get(qn("w:eastAsia")),
        "cs": r_fonts.get(qn("w:cs")),
    }


def test_normal_style_defines_latin_and_east_asian_fonts() -> None:
    doc = Document()

    _setup_document_styles(doc)

    rpr = doc.styles["Normal"]._element.rPr
    assert _font_attributes(rpr) == {
        "ascii": LATIN_FONT,
        "hAnsi": LATIN_FONT,
        "eastAsia": EAST_ASIA_FONT,
        "cs": LATIN_FONT,
    }

    language = rpr.find(qn("w:lang"))
    assert language.get(qn("w:val")) == DOCUMENT_LANGUAGE
    assert language.get(qn("w:eastAsia")) == EAST_ASIA_LANGUAGE


def test_inline_code_uses_monospace_latin_and_document_cjk_font() -> None:
    doc = Document()
    _setup_document_styles(doc)
    paragraph = doc.add_paragraph()

    _add_inline_formatting(paragraph, "`配置 config`")

    assert len(paragraph.runs) == 1
    assert _font_attributes(paragraph.runs[0]._element.rPr) == {
        "ascii": MONOSPACE_FONT,
        "hAnsi": MONOSPACE_FONT,
        "eastAsia": EAST_ASIA_FONT,
        "cs": MONOSPACE_FONT,
    }
