# SPDX-FileCopyrightText: 2025 Weibo, Inc.
#
# SPDX-License-Identifier: Apache-2.0

"""
Unit tests for docx_generator prompt sanitization.

Verifies that system-injected metadata blocks (<system-reminder>) stored in
Subtask.prompt are stripped before being written to the exported DOCX document.
"""

import json
from datetime import datetime
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from app.schemas.subtask import SubtaskRole
from app.services.export.docx_generator import _add_message


def _make_user_subtask(prompt: str, contexts=None):
    """Create a minimal mock Subtask with USER role."""
    subtask = MagicMock()
    subtask.role = MagicMock()
    subtask.role.value = "USER"
    subtask.prompt = prompt
    subtask.result = None
    subtask.contexts = contexts or []
    subtask.sender_user_id = None
    subtask.updated_at = datetime(2025, 1, 1, 12, 0, 0)
    return subtask


def _make_assistant_subtask(result_value: str):
    """Create a minimal mock Subtask with ASSISTANT role."""
    subtask = MagicMock()
    subtask.role = MagicMock()
    subtask.role.value = "ASSISTANT"
    subtask.prompt = None
    subtask.result = {"value": result_value}
    subtask.contexts = []
    subtask.sender_user_id = None
    subtask.updated_at = datetime(2025, 1, 1, 12, 0, 0)
    return subtask


def _make_task_and_user():
    """Create mock task (Kind) and user objects."""
    task = MagicMock()
    task.user_id = 1
    task.json = {"spec": {"teamRef": {"name": "TestTeam"}}}
    user = MagicMock()
    user.user_name = "Alice"
    return task, user


class TestDocxGeneratorPromptSanitization:
    """Tests for prompt sanitization in _add_message."""

    def test_plain_text_prompt_written_to_document(self):
        """Plain text user prompt appears as-is in DOCX content."""
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_user_subtask("What is the capital of France?")

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        # Collect all paragraph text from the document
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "What is the capital of France?" in full_text

    def test_system_reminder_stripped_from_docx_output(self):
        """<system-reminder> block in JSON array prompt must NOT appear in DOCX."""
        raw_prompt = json.dumps(
            [
                {"type": "text", "text": "Summarize the attached report"},
                {
                    "type": "text",
                    "text": "<system-reminder><CurrentTime>2025-01-01 09:00</CurrentTime></system-reminder>",
                },
            ]
        )
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_user_subtask(raw_prompt)

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        full_text = "\n".join(p.text for p in doc.paragraphs)

        # User text must be present
        assert "Summarize the attached report" in full_text
        # System metadata must NOT appear
        assert "<system-reminder>" not in full_text
        assert "CurrentTime" not in full_text
        # Raw JSON array bracket must NOT appear as literal text
        assert full_text.count("[{") == 0

    def test_json_array_with_multiple_system_blocks(self):
        """All extra system blocks are removed; only user text is written."""
        raw_prompt = json.dumps(
            [
                {"type": "text", "text": "Review this codebase"},
                {
                    "type": "text",
                    "text": "<system-reminder><Attachment>repo.zip</Attachment></system-reminder>",
                },
                {
                    "type": "text",
                    "text": "<system-reminder><CurrentTime>2025-03-01</CurrentTime></system-reminder>",
                },
            ]
        )
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_user_subtask(raw_prompt)

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Review this codebase" in full_text
        assert "<system-reminder>" not in full_text

    def test_assistant_message_uses_result_not_prompt(self):
        """ASSISTANT messages render result.value, not prompt."""
        doc = Document()
        task, user = _make_task_and_user()
        subtask = _make_assistant_subtask("Paris is the capital of France.")

        db = MagicMock()
        db.query.return_value.filter.return_value.first.return_value = user

        _add_message(doc, subtask, task, user, db)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Paris is the capital of France." in full_text
