# SPDX-FileCopyrightText: 2025 Weibo, Inc.
#
# SPDX-License-Identifier: Apache-2.0

"""
Storage service for evaluation module.

Handles S3/MinIO file storage for questions, criteria, answers, and reports.
"""

import logging
import os
from datetime import datetime, timedelta
from io import BytesIO
from typing import Optional
from urllib.parse import quote

from minio import Minio
from minio.error import S3Error

from app.core.config import settings

logger = logging.getLogger(__name__)


class EvalStorageService:
    """
    Storage service for evaluation module files.

    Uses MinIO/S3 for storing:
    - Question content attachments
    - Grading criteria attachments
    - Answer attachments
    - Grading reports
    """

    def __init__(self):
        """Initialize storage service with MinIO client."""
        self._client: Optional[Minio] = None
        self._bucket = settings.ATTACHMENT_S3_BUCKET
        self._prefix = os.getenv("EVAL_S3_PREFIX", "evaluation")
        self._presigned_expires = int(
            os.getenv("GRADING_PRESIGNED_URL_EXPIRES", "3600")
        )

    @property
    def client(self) -> Optional[Minio]:
        """Lazy initialization of MinIO client."""
        if self._client is not None:
            return self._client

        # Check if MinIO/S3 is configured
        if settings.ATTACHMENT_STORAGE_BACKEND not in ("minio", "s3"):
            logger.warning(
                "[Evaluation] MinIO/S3 storage not configured for evaluation module"
            )
            return None

        if not settings.ATTACHMENT_S3_ENDPOINT:
            logger.warning("[Evaluation] ATTACHMENT_S3_ENDPOINT not configured")
            return None

        try:
            endpoint = settings.ATTACHMENT_S3_ENDPOINT
            use_ssl = settings.ATTACHMENT_S3_USE_SSL

            # Remove protocol prefix if present
            if endpoint.startswith("http://"):
                endpoint = endpoint[7:]
                use_ssl = False
            elif endpoint.startswith("https://"):
                endpoint = endpoint[8:]
                use_ssl = True

            self._client = Minio(
                endpoint,
                access_key=settings.ATTACHMENT_S3_ACCESS_KEY,
                secret_key=settings.ATTACHMENT_S3_SECRET_KEY,
                secure=use_ssl,
                region=settings.ATTACHMENT_S3_REGION,
            )

            # Ensure bucket exists
            if not self._client.bucket_exists(self._bucket):
                self._client.make_bucket(
                    self._bucket, location=settings.ATTACHMENT_S3_REGION
                )
                logger.info(f"[Evaluation] Created bucket: {self._bucket}")

            return self._client

        except Exception as e:
            logger.error(f"[Evaluation] Failed to initialize MinIO client: {e}")
            return None

    def _build_key(self, category: str, *parts: str) -> str:
        """
        Build storage key from parts.

        Args:
            category: Category (questions, criteria, answers, reports)
            *parts: Path parts

        Returns:
            Full storage key
        """
        return "/".join([self._prefix, category, *parts])

    def upload_question_content(
        self,
        topic_id: int,
        question_id: int,
        version: str,
        filename: str,
        data: bytes,
        content_type: str = "application/octet-stream",
    ) -> Optional[str]:
        """
        Upload question content attachment.

        Args:
            topic_id: Topic ID
            question_id: Question ID
            version: Question version
            filename: Original filename
            data: File data
            content_type: MIME type

        Returns:
            Storage key if successful
        """
        if not self.client:
            return None

        key = self._build_key(
            "questions",
            str(topic_id),
            str(question_id),
            version,
            "content",
            filename,
        )

        return self._upload(key, data, content_type, filename)

    def upload_criteria(
        self,
        topic_id: int,
        question_id: int,
        version: str,
        filename: str,
        data: bytes,
        content_type: str = "application/octet-stream",
    ) -> Optional[str]:
        """
        Upload grading criteria attachment.

        Args:
            topic_id: Topic ID
            question_id: Question ID
            version: Question version
            filename: Original filename
            data: File data
            content_type: MIME type

        Returns:
            Storage key if successful
        """
        if not self.client:
            return None

        key = self._build_key(
            "criteria",
            str(topic_id),
            str(question_id),
            version,
            filename,
        )

        return self._upload(key, data, content_type, filename)

    def upload_answer(
        self,
        respondent_id: int,
        topic_id: int,
        question_id: int,
        filename: str,
        data: bytes,
        content_type: str = "application/octet-stream",
    ) -> Optional[str]:
        """
        Upload answer attachment.

        Args:
            respondent_id: Respondent user ID
            topic_id: Topic ID
            question_id: Question ID
            filename: Original filename
            data: File data
            content_type: MIME type

        Returns:
            Storage key if successful
        """
        if not self.client:
            return None

        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        key = self._build_key(
            "answers",
            str(respondent_id),
            str(topic_id),
            str(question_id),
            timestamp,
            filename,
        )

        return self._upload(key, data, content_type, filename)

    def save_grading_report(
        self,
        respondent_id: int,
        topic_id: int,
        question_id: int,
        content: str,
        is_draft: bool = True,
    ) -> Optional[str]:
        """
        Save grading report to S3.

        Args:
            respondent_id: Respondent user ID
            topic_id: Topic ID
            question_id: Question ID
            content: Report content (Markdown)
            is_draft: Whether this is a draft or final report

        Returns:
            Storage key if successful
        """
        if not self.client:
            return None

        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
        filename = "draft.md" if is_draft else "final.md"
        key = self._build_key(
            "reports",
            str(respondent_id),
            str(topic_id),
            str(question_id),
            timestamp,
            filename,
        )

        data = content.encode("utf-8")
        return self._upload(key, data, "text/markdown", filename)

    def _upload(
        self,
        key: str,
        data: bytes,
        content_type: str,
        filename: str,
    ) -> Optional[str]:
        """
        Upload data to S3.

        Args:
            key: Storage key
            data: File data
            content_type: MIME type
            filename: Original filename

        Returns:
            Storage key if successful
        """
        try:
            # URL-encode filename for metadata
            encoded_filename = quote(filename, safe="")

            self.client.put_object(
                self._bucket,
                key,
                BytesIO(data),
                length=len(data),
                content_type=content_type,
                metadata={"filename": encoded_filename},
            )

            logger.info(f"[Evaluation] Uploaded file to S3: {key}")
            return key

        except S3Error as e:
            logger.error(f"[Evaluation] Failed to upload to S3: {e}")
            return None
        except Exception as e:
            logger.error(f"[Evaluation] Unexpected error uploading to S3: {e}")
            return None

    def get(self, key: str) -> Optional[bytes]:
        """
        Get file data from S3.

        Args:
            key: Storage key

        Returns:
            File data if found
        """
        if not self.client:
            return None

        try:
            response = self.client.get_object(self._bucket, key)
            data = response.read()
            response.close()
            response.release_conn()
            return data

        except S3Error as e:
            if e.code == "NoSuchKey":
                logger.info(f"[Evaluation] Object not found: {key}")
            else:
                logger.error(f"[Evaluation] Failed to get from S3: {e}")
            return None
        except Exception as e:
            logger.error(f"[Evaluation] Unexpected error getting from S3: {e}")
            return None

    def delete(self, key: str) -> bool:
        """
        Delete file from S3.

        Args:
            key: Storage key

        Returns:
            True if deleted
        """
        if not self.client:
            return False

        try:
            self.client.remove_object(self._bucket, key)
            logger.info(f"[Evaluation] Deleted file from S3: {key}")
            return True

        except S3Error as e:
            if e.code == "NoSuchKey":
                return True
            logger.error(f"[Evaluation] Failed to delete from S3: {e}")
            return False
        except Exception as e:
            logger.error(f"[Evaluation] Unexpected error deleting from S3: {e}")
            return False

    def exists(self, key: str) -> bool:
        """
        Check if file exists in S3.

        Args:
            key: Storage key

        Returns:
            True if exists
        """
        if not self.client:
            return False

        try:
            self.client.stat_object(self._bucket, key)
            return True

        except S3Error as e:
            if e.code == "NoSuchKey":
                return False
            logger.error(f"[Evaluation] Failed to check existence in S3: {e}")
            return False
        except Exception as e:
            logger.error(f"[Evaluation] Unexpected error checking existence in S3: {e}")
            return False

    def get_presigned_url(
        self, key: str, expires: Optional[int] = None, force_download: bool = True
    ) -> Optional[str]:
        """
        Generate presigned URL for file access (GET).

        Args:
            key: Storage key
            expires: Expiration time in seconds (default from config)
            force_download: If True, set Content-Disposition header to force download

        Returns:
            Presigned URL if successful
        """
        if not self.client:
            return None

        if expires is None:
            expires = self._presigned_expires

        try:
            # Build response headers for download
            response_headers = {}
            if force_download:
                # Extract filename from key (last part of the path)
                filename = key.split("/")[-1]
                # RFC 5987 encoding for non-ASCII filenames
                # Use both filename for ASCII clients and filename* for modern browsers
                encoded_filename = quote(filename, safe="")
                response_headers["response-content-disposition"] = (
                    f'attachment; filename="{encoded_filename}"; '
                    f"filename*=UTF-8''{encoded_filename}"
                )

            url = self.client.presigned_get_object(
                self._bucket,
                key,
                expires=timedelta(seconds=expires),
                response_headers=response_headers if response_headers else None,
            )
            return url

        except S3Error as e:
            logger.error(f"[Evaluation] Failed to generate presigned GET URL: {e}")
            return None
        except Exception as e:
            logger.error(
                f"[Evaluation] Unexpected error generating presigned GET URL: {e}"
            )
            return None

    def get_presigned_put_url(
        self, key: str, expires: Optional[int] = None
    ) -> Optional[str]:
        """
        Generate presigned URL for file upload (PUT).

        Args:
            key: Storage key
            expires: Expiration time in seconds (default from config)

        Returns:
            Presigned PUT URL if successful
        """
        if not self.client:
            return None

        if expires is None:
            expires = self._presigned_expires

        try:
            url = self.client.presigned_put_object(
                self._bucket,
                key,
                expires=timedelta(seconds=expires),
            )
            return url

        except S3Error as e:
            logger.error(f"[Evaluation] Failed to generate presigned PUT URL: {e}")
            return None
        except Exception as e:
            logger.error(
                f"[Evaluation] Unexpected error generating presigned PUT URL: {e}"
            )
            return None

    def generate_upload_key(
        self,
        file_type: str,
        user_id: int,
        topic_id: int,
        question_id: Optional[int] = None,
        filename: str = "",
    ) -> str:
        """
        Generate a storage key for file upload based on file type.

        Args:
            file_type: Type of file (question_content, question_criteria, answer_attachment)
            user_id: User ID uploading the file
            topic_id: Topic ID
            question_id: Question ID (optional, depends on type)
            filename: Original filename

        Returns:
            Generated storage key
        """
        timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")

        if file_type == "question_content":
            return self._build_key(
                "questions",
                str(topic_id),
                str(question_id or 0),
                "draft",
                "content",
                f"{timestamp}_{filename}",
            )
        elif file_type == "question_criteria":
            return self._build_key(
                "criteria",
                str(topic_id),
                str(question_id or 0),
                "draft",
                f"{timestamp}_{filename}",
            )
        elif file_type == "answer_attachment":
            return self._build_key(
                "answers",
                str(user_id),
                str(topic_id),
                str(question_id or 0),
                timestamp,
                filename,
            )
        else:
            # Default fallback
            return self._build_key(
                "uploads",
                str(user_id),
                timestamp,
                filename,
            )

    def extract_text_from_file(self, key: str) -> Optional[str]:
        """
        Extract text content from a file stored in S3.

        Supports:
        - Plain text files (.txt, .md)
        - Word documents (.docx)
        - PDF files (.pdf)
        - Other binary files (returns a notice that content cannot be extracted)

        Args:
            key: Storage key in S3

        Returns:
            Extracted text content, or None if extraction fails
        """
        # Get file data from S3
        data = self.get(key)
        if not data:
            logger.warning(f"[Evaluation] Could not get file from S3: {key}")
            return None

        # Determine file type from key (filename)
        filename = key.split("/")[-1].lower()

        try:
            if filename.endswith((".txt", ".md", ".markdown")):
                return self._extract_text_plain(data)
            elif filename.endswith(".docx"):
                return self._extract_text_docx(data)
            elif filename.endswith(".pdf"):
                return self._extract_text_pdf(data)
            else:
                # For unsupported formats, return a notice
                return f"[File: {filename} - content extraction not supported for this format]"

        except Exception as e:
            logger.error(f"[Evaluation] Failed to extract text from {key}: {e}")
            return f"[File: {filename} - failed to extract content: {str(e)[:100]}]"

    def _extract_text_plain(self, data: bytes) -> str:
        """Extract text from plain text file."""
        # Try different encodings
        for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        # Fallback with replacement
        return data.decode("utf-8", errors="replace")

    def _extract_text_docx(self, data: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            from io import BytesIO

            from docx import Document

            doc = Document(BytesIO(data))
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            return "\n\n".join(paragraphs)

        except ImportError:
            logger.warning(
                "[Evaluation] python-docx not installed, cannot extract DOCX"
            )
            return "[DOCX file - python-docx library not available for extraction]"
        except Exception as e:
            logger.error(f"[Evaluation] Failed to extract DOCX: {e}")
            return f"[DOCX file - extraction failed: {str(e)[:100]}]"

    def _extract_text_pdf(self, data: bytes) -> str:
        """Extract text from PDF file."""
        try:
            from io import BytesIO

            from PyPDF2 import PdfReader

            reader = PdfReader(BytesIO(data))
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{page_text.strip()}")

            return (
                "\n\n".join(text_parts)
                if text_parts
                else "[PDF file - no extractable text]"
            )

        except ImportError:
            logger.warning("[Evaluation] PyPDF2 not installed, cannot extract PDF")
            return "[PDF file - PyPDF2 library not available for extraction]"
        except Exception as e:
            logger.error(f"[Evaluation] Failed to extract PDF: {e}")
            return f"[PDF file - extraction failed: {str(e)[:100]}]"
