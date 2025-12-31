# SPDX-FileCopyrightText: 2025 Weibo, Inc.
#
# SPDX-License-Identifier: Apache-2.0

"""
DOCX document parser.

Parses Microsoft Word documents into structured blocks.
"""

import io
import logging
import uuid
from typing import List, Optional

from app.services.document_parser.base import BaseParser
from app.services.document_parser.factory import ParserFactory
from app.services.document_parser.models.block import BlockType, DocumentBlockData, SourceType

logger = logging.getLogger(__name__)


@ParserFactory.register
class DocxParser(BaseParser):
    """
    Parser for Microsoft Word (.docx) documents.

    Uses python-docx to extract paragraphs, headings, tables, and images.
    Text blocks are non-editable.
    """

    # Mapping of Word heading styles to levels
    HEADING_STYLES = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Heading 5": 5,
        "Heading 6": 6,
        "Title": 1,
        "Subtitle": 2,
    }

    def supported_content_types(self) -> List[str]:
        """Return supported MIME types."""
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]

    def supported_extensions(self) -> List[str]:
        """Return supported file extensions."""
        return [".docx"]

    def parse(
        self,
        binary_data: bytes,
        document_id: str,
        filename: str,
    ) -> List[DocumentBlockData]:
        """
        Parse DOCX content into blocks.

        Args:
            binary_data: Raw DOCX file content
            document_id: Document identifier
            filename: Original filename

        Returns:
            List of DocumentBlockData blocks
        """
        logger.info(f"Parsing DOCX document: {filename}")

        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return [
                DocumentBlockData(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    source_type=SourceType.DOCX,
                    block_type=BlockType.UNSUPPORTED,
                    content="DOCX parsing requires python-docx library. Please install it with: pip install python-docx",
                    editable=False,
                    order_index=0,
                )
            ]

        blocks: List[DocumentBlockData] = []
        order_index = 0

        try:
            # Open document from bytes
            doc = Document(io.BytesIO(binary_data))

            # Process paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if not text:
                    continue

                style_name = paragraph.style.name if paragraph.style else ""

                # Check if this is a heading
                if style_name in self.HEADING_STYLES:
                    blocks.append(
                        DocumentBlockData(
                            id=str(uuid.uuid4()),
                            document_id=document_id,
                            source_type=SourceType.DOCX,
                            block_type=BlockType.HEADING,
                            content=text,
                            editable=False,
                            order_index=order_index,
                            source_ref={"paragraph_index": para_idx},
                            metadata={"level": self.HEADING_STYLES[style_name]},
                        )
                    )
                # Check for list items
                elif self._is_list_paragraph(paragraph):
                    # Collect consecutive list items
                    blocks.append(
                        DocumentBlockData(
                            id=str(uuid.uuid4()),
                            document_id=document_id,
                            source_type=SourceType.DOCX,
                            block_type=BlockType.LIST,
                            content=text,
                            editable=False,
                            order_index=order_index,
                            source_ref={"paragraph_index": para_idx},
                        )
                    )
                else:
                    blocks.append(
                        DocumentBlockData(
                            id=str(uuid.uuid4()),
                            document_id=document_id,
                            source_type=SourceType.DOCX,
                            block_type=BlockType.PARAGRAPH,
                            content=text,
                            editable=False,
                            order_index=order_index,
                            source_ref={"paragraph_index": para_idx},
                        )
                    )

                order_index += 1

            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_content = self._extract_table_content(table)
                if table_content:
                    blocks.append(
                        DocumentBlockData(
                            id=str(uuid.uuid4()),
                            document_id=document_id,
                            source_type=SourceType.DOCX,
                            block_type=BlockType.TABLE,
                            content=table_content,
                            editable=False,
                            order_index=order_index,
                            source_ref={"table_index": table_idx},
                        )
                    )
                    order_index += 1

            # Process images from relationships
            image_index = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image_data = rel.target_part.blob
                        image_ext = rel.target_part.content_type.split("/")[-1]

                        # Save image
                        image_url = None
                        if self.storage:
                            image_filename = f"image_{image_index}.{image_ext}"
                            try:
                                image_url = self.storage.save_image_sync(
                                    image_data,
                                    image_filename,
                                    document_id,
                                    rel.target_part.content_type,
                                )
                            except Exception as e:
                                logger.warning(f"Failed to save image: {e}")

                        # Get image description
                        description = None
                        ocr_text = None
                        if self.ocr:
                            try:
                                description = self.ocr.describe_image_sync(image_data)
                                ocr_text = self.ocr.extract_text_sync(image_data)
                            except Exception as e:
                                logger.warning(f"Failed to process image with OCR: {e}")

                        blocks.append(
                            DocumentBlockData(
                                id=str(uuid.uuid4()),
                                document_id=document_id,
                                source_type=SourceType.DOCX,
                                block_type=BlockType.IMAGE,
                                content=description or "[Image - no description available]",
                                editable=False,
                                order_index=order_index,
                                source_ref={"image_index": image_index},
                                metadata={
                                    "image_url": image_url,
                                    "ocr_text": ocr_text,
                                },
                            )
                        )
                        order_index += 1
                        image_index += 1

                    except Exception as e:
                        logger.warning(f"Failed to extract image: {e}")

            logger.info(f"Parsed {len(blocks)} blocks from DOCX document")

        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            blocks.append(
                DocumentBlockData(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    source_type=SourceType.DOCX,
                    block_type=BlockType.UNSUPPORTED,
                    content=f"Error parsing DOCX: {str(e)}",
                    editable=False,
                    order_index=0,
                )
            )

        return blocks

    def _is_list_paragraph(self, paragraph) -> bool:
        """
        Check if a paragraph is a list item.

        Args:
            paragraph: python-docx Paragraph object

        Returns:
            True if paragraph is a list item
        """
        try:
            # Check for numbering properties
            p_pr = paragraph._element.pPr
            if p_pr is not None:
                num_pr = p_pr.numPr
                if num_pr is not None:
                    return True

            # Check style name for list indicators
            style_name = paragraph.style.name if paragraph.style else ""
            if any(
                keyword in style_name.lower()
                for keyword in ["list", "bullet", "number"]
            ):
                return True

        except Exception:
            pass

        return False

    def _extract_table_content(self, table) -> str:
        """
        Extract table content as Markdown format.

        Args:
            table: python-docx Table object

        Returns:
            Table content in Markdown format
        """
        rows = []

        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

            # Add separator after header row
            if row_idx == 0:
                separator = "|" + "|".join([" --- "] * len(cells)) + "|"
                rows.append(separator)

        return "\n".join(rows) if rows else ""
