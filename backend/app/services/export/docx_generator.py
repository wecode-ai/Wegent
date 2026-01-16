# SPDX-FileCopyrightText: 2025 Weibo, Inc.
#
# SPDX-License-Identifier: Apache-2.0

"""
DOCX document generator for task export.
Generates formatted Word documents with markdown rendering and image embedding.
"""

import io
import logging
import re
from datetime import datetime
from typing import BinaryIO, List, Optional

import emoji
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

from app.models.kind import Kind
from app.models.subtask import Subtask
from app.models.subtask_context import SubtaskContext
from app.models.user import User

logger = logging.getLogger(__name__)

# Wegent brand color (Mint Blue)
PRIMARY_COLOR = RGBColor(20, 184, 166)  # #14B8A6
TEXT_COLOR = RGBColor(36, 41, 46)
CODE_BG_COLOR = RGBColor(246, 248, 250)
LINK_COLOR = RGBColor(85, 185, 247)


def generate_task_docx(
    task: Kind, db: Session, message_ids: Optional[List[int]] = None
) -> io.BytesIO:
    """
    Generate a DOCX document from task data.

    Args:
        task: Task Kind instance
        db: Database session
        message_ids: Optional list of subtask IDs to include. If None, includes all subtasks.

    Returns:
        BytesIO buffer containing the DOCX document
    """
    doc = Document()

    # Set document properties
    task_data = task.json.get("spec", {})
    task_title = (
        task.json.get("metadata", {}).get("name", "")
        or task_data.get("title", "")
        or task_data.get("prompt", "Chat Export")[:50]
    )

    doc.core_properties.title = task_title
    doc.core_properties.author = "Wegent AI"
    doc.core_properties.created = datetime.now()

    # Configure document styles
    _setup_document_styles(doc)

    # Add header
    _add_document_header(doc, task_title)

    # Add content
    _add_task_content(doc, task, db, message_ids=message_ids)

    # Add footer
    _add_document_footer(doc)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def _setup_document_styles(doc: Document):
    """Configure document-wide styles with emoji support for Word, WPS, and macOS"""
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    # Use Arial which has better emoji support than Calibri
    font.name = "Arial"
    font.size = Pt(11)
    font.color.rgb = TEXT_COLOR

    # Note: We don't set document-wide emoji font here because different platforms
    # have different emoji fonts (Apple Color Emoji on macOS, Segoe UI Emoji on Windows).
    # Instead, we rely on the Office application's automatic font fallback mechanism,
    # which works well in modern versions (Word 2016+, macOS Word, WPS Office).
    #
    # The key is to ensure text is properly UTF-8 encoded, which python-docx handles automatically.

    # Set paragraph spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(12)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _add_document_header(doc: Document, task_title: str):
    """Add document header with logo and title"""
    # Logo text
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo.add_run("Wegent AI")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    # Task title (support emoji)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text_with_emoji_support(title, task_title, bold=True)
    # Set title font size for all runs
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = PRIMARY_COLOR

    # Divider
    _add_horizontal_rule(doc)
    doc.add_paragraph()  # Spacing


def _add_horizontal_rule(doc: Document):
    """Add horizontal line separator"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "14B8A6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_task_content(
    doc: Document, task: Kind, db: Session, message_ids: Optional[List[int]] = None
):
    """Add task subtasks as messages, optionally filtered by message_ids"""
    # Query subtasks with attachments
    query = db.query(Subtask).filter(Subtask.task_id == task.id)

    # Filter by message_ids if provided
    if message_ids is not None and len(message_ids) > 0:
        query = query.filter(Subtask.id.in_(message_ids))

    # Order by id to maintain original message order
    subtasks = query.order_by(Subtask.id.asc()).all()

    # Get user for display name
    user = db.query(User).filter(User.id == task.user_id).first()

    for subtask in subtasks:
        _add_message(doc, subtask, task, user, db)


def _add_message(doc: Document, subtask: Subtask, task: Kind, user: User, db: Session):
    """Add a single message (user or AI)"""
    is_user = subtask.role.value == "USER"

    # Message header (sender + timestamp)
    header = doc.add_paragraph()

    # Sender name (support emoji in names)
    if is_user:
        # For group chat messages, check sender_user_id
        if subtask.sender_user_id and subtask.sender_user_id > 0:
            # Query the actual sender from database
            actual_sender = (
                db.query(User).filter(User.id == subtask.sender_user_id).first()
            )
            sender_name = actual_sender.user_name if actual_sender else "User"
        else:
            # Regular message from task owner
            sender_name = user.user_name if user else "User"
    else:
        # Get team name from task
        task_data = task.json.get("spec", {})
        team_ref = task_data.get("teamRef", {})
        sender_name = team_ref.get("name", "AI Assistant")

    _add_text_with_emoji_support(header, f"{sender_name}: ", bold=True)
    # Set sender font properties for all runs just added
    for run in header.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR if is_user else PRIMARY_COLOR

    # Timestamp
    timestamp = subtask.updated_at.strftime("%Y-%m-%d %H:%M:%S")
    time_run = header.add_run(timestamp)
    time_run.font.size = Pt(9)
    time_run.font.color.rgb = RGBColor(160, 160, 160)

    # Contexts (attachments and knowledge bases) for user messages
    if is_user and subtask.contexts:
        from app.models.subtask_context import ContextType

        # Filter attachment type contexts
        attachment_contexts = [
            ctx
            for ctx in subtask.contexts
            if ctx.context_type == ContextType.ATTACHMENT.value
        ]
        if attachment_contexts:
            _add_attachments(doc, attachment_contexts, db)

        # Filter knowledge base type contexts
        knowledge_base_contexts = [
            ctx
            for ctx in subtask.contexts
            if ctx.context_type == ContextType.KNOWLEDGE_BASE.value
        ]
        if knowledge_base_contexts:
            _add_knowledge_bases(doc, knowledge_base_contexts)

    # Message content
    content = subtask.prompt if is_user else _extract_result_value(subtask.result)

    # Remove special markers
    content = _clean_content(content)

    # Convert emoji to text
    content = _convert_emoji_to_text(content)

    # Parse and render markdown
    _render_markdown_content(doc, content)

    # Add spacing after message
    doc.add_paragraph()


def _extract_result_value(result: any) -> str:
    """Extract text content from subtask result"""
    if not result:
        return ""

    if isinstance(result, str):
        return result

    if isinstance(result, dict):
        if "value" in result and result["value"]:
            return str(result["value"])
        if "thinking" in result:
            return ""  # Don't show thinking in export
        return str(result)

    return str(result)


def _clean_content(content: str) -> str:
    """Remove special markers and formatting"""
    if not content:
        return ""

    # Remove special markers
    content = re.sub(r"\$\{\$\$\}\$", "\n", content)
    content = re.sub(r"__PROGRESS_BAR__:.*?:\d+", "", content)
    content = re.sub(r"__PROMPT_TRUNCATED__:.*?::(.*?)(?=\n|$)", r"\1", content)

    return content.strip()


def _convert_emoji_to_text(text: str) -> str:
    """
    Keep emoji as-is for Word document.
    Modern Word versions (2016+) support Unicode emoji natively.
    The key is to ensure proper UTF-8 encoding when adding text to the document.
    """
    if not text:
        return text

    # Keep the original text with emoji
    # python-docx handles Unicode correctly if we don't interfere
    return text


def _add_attachments(doc: Document, attachments: List[SubtaskContext], db: Session):
    """Add attachment section (images embedded, files as info cards)"""
    for attachment in attachments:
        # Get file extension from type_data
        type_data = attachment.type_data or {}
        file_extension = type_data.get("file_extension", "")
        if _is_image_extension(file_extension):
            _add_image_attachment(doc, attachment, db)
        else:
            _add_file_attachment(doc, attachment)


def _is_image_extension(extension: str) -> bool:
    """Check if file extension is an image type"""
    image_exts = [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"]
    return extension.lower() in image_exts


def _add_image_attachment(doc: Document, attachment: SubtaskContext, db: Session):
    """Embed image attachment in document"""
    try:
        from app.services.attachment.storage_factory import get_storage_backend

        # Get storage backend
        storage_backend = get_storage_backend(db)

        # Get storage info from type_data
        type_data = attachment.type_data or {}
        storage_backend_type = type_data.get("storage_backend", "mysql")

        # Get image data based on storage backend
        if storage_backend_type == "mysql":
            image_data = attachment.binary_data
        else:
            # External storage (S3/MinIO)
            storage_key = type_data.get("storage_key", "")
            image_data = storage_backend.get(storage_key)

        if image_data:
            # Save to temporary buffer
            image_buffer = io.BytesIO(image_data)

            # Add image to document (max width 6 inches)
            doc.add_picture(image_buffer, width=Inches(6))

            # Add caption
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(attachment.name)
            run.font.size = Pt(8)
            run.font.italic = True
            run.font.color.rgb = RGBColor(150, 150, 150)
        else:
            _add_file_attachment(doc, attachment)

    except Exception as e:
        logger.warning(f"Failed to embed image {attachment.id}: {e}")
        _add_file_attachment(doc, attachment)


def _add_file_attachment(doc: Document, attachment: SubtaskContext):
    """Add file attachment info card"""
    p = doc.add_paragraph()

    # Get file info from type_data
    type_data = attachment.type_data or {}
    file_extension = type_data.get("file_extension", "")
    file_size = type_data.get("file_size", 0)

    # File type label
    file_type = _get_file_type_label(file_extension)
    type_run = p.add_run(f"{file_type} ")
    type_run.font.bold = True
    type_run.font.size = Pt(9)
    type_run.font.color.rgb = RGBColor(100, 100, 100)

    # Filename
    name_run = p.add_run(attachment.name)
    name_run.font.size = Pt(10)

    # File size
    size_run = p.add_run(f" ({_format_file_size(file_size)})")
    size_run.font.size = Pt(9)
    size_run.font.color.rgb = RGBColor(150, 150, 150)

    # Add background shading
    p_fmt = p.paragraph_format
    p_fmt.left_indent = Inches(0.2)
    p_fmt.space_after = Pt(6)


def _add_knowledge_bases(doc: Document, knowledge_bases: List[SubtaskContext]):
    """Add knowledge base info cards"""
    for kb in knowledge_bases:
        p = doc.add_paragraph()

        # Get knowledge base info from type_data
        type_data = kb.type_data or {}
        document_count = type_data.get("document_count", 0)

        # Knowledge base type label (gray color, consistent with attachment labels)
        type_run = p.add_run("[KB] ")
        type_run.font.bold = True
        type_run.font.size = Pt(9)
        type_run.font.color.rgb = RGBColor(100, 100, 100)

        # Knowledge base name
        name_run = p.add_run(kb.name)
        name_run.font.size = Pt(10)

        # Document count
        count_run = p.add_run(f" ({document_count} docs)")
        count_run.font.size = Pt(9)
        count_run.font.color.rgb = RGBColor(150, 150, 150)

        # Add background shading
        p_fmt = p.paragraph_format
        p_fmt.left_indent = Inches(0.2)
        p_fmt.space_after = Pt(6)


def _get_file_type_label(extension: str) -> str:
    """Get file type label from extension"""
    ext_map = {
        ".pdf": "[PDF]",
        ".doc": "[DOC]",
        ".docx": "[DOC]",
        ".txt": "[TXT]",
        ".md": "[MD]",
        ".zip": "[ZIP]",
        ".rar": "[RAR]",
        ".xls": "[XLS]",
        ".xlsx": "[XLSX]",
        ".ppt": "[PPT]",
        ".pptx": "[PPT]",
    }
    return ext_map.get(extension.lower(), "[FILE]")


def _format_file_size(size_bytes: int) -> str:
    """Format file size in human-readable format"""
    for unit in ["B", "KB", "MB", "GB"]:
        if size_bytes < 1024.0:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024.0
    return f"{size_bytes:.1f} TB"


def _render_markdown_content(doc: Document, content: str):
    """
    Parse and render markdown content.
    Supports: headings, bold, italic, code, code blocks, lists, quotes, tables, links.
    """
    if not content:
        return

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    code_language = ""

    # Table state
    in_table = False
    table_headers = []
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block delimiter
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_language = line.strip()[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                _add_code_block(doc, "\n".join(code_lines), code_language)
                code_lines = []
                code_language = ""
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Parse line type
        line_type, parsed_data = _parse_line_type(line)

        # Handle table
        if line_type == "table_separator":
            # Next lines are table rows
            in_table = True
            if i > 0:
                # Previous line was header
                header_type, header_data = _parse_line_type(lines[i - 1])
                if header_type == "table_row":
                    table_headers = header_data["cells"]
            i += 1
            continue

        if line_type == "table_row" and in_table:
            table_rows.append(parsed_data["cells"])
            i += 1
            continue

        # Flush table if we left table context
        if in_table and line_type != "table_row":
            _add_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

        # Render line
        _render_markdown_line(doc, line_type, parsed_data)
        i += 1

    # Flush remaining table
    if in_table and table_headers:
        _add_table(doc, table_headers, table_rows)


def _parse_line_type(line: str) -> tuple:
    """Parse markdown line type and extract data"""
    stripped = line.strip()

    if not stripped:
        return ("empty", {})

    # Heading
    heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if heading_match:
        level = len(heading_match.group(1))
        content = heading_match.group(2)
        return ("heading", {"level": level, "content": content})

    # Unordered list
    if re.match(r"^[-*+]\s+", stripped):
        content = re.sub(r"^[-*+]\s+", "", stripped)
        return ("list_unordered", {"content": content})

    # Ordered list
    ordered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
    if ordered_match:
        return (
            "list_ordered",
            {"number": ordered_match.group(1), "content": ordered_match.group(2)},
        )

    # Blockquote
    if stripped.startswith(">"):
        content = re.sub(r"^>\s*", "", stripped)
        return ("blockquote", {"content": content})

    # Horizontal rule
    if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
        return ("hr", {})

    # Table separator
    if re.match(r"^\|?[\s]*:?-+:?[\s]*(\|[\s]*:?-+:?[\s]*)+\|?$", stripped):
        return ("table_separator", {})

    # Table row
    if "|" in stripped:
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if cells:
            return ("table_row", {"cells": cells})

    # Regular paragraph
    return ("paragraph", {"content": stripped})


def _render_markdown_line(doc: Document, line_type: str, data: dict):
    """Render a single markdown line"""
    if line_type == "empty":
        doc.add_paragraph()
        return

    if line_type == "heading":
        level = data["level"]
        content = data["content"]
        p = doc.add_paragraph()
        _add_inline_formatting(p, content)
        # Apply heading formatting to all runs
        for run in p.runs:
            run.font.size = Pt(18 - level * 2)
            run.font.bold = True
            run.font.color.rgb = TEXT_COLOR
        if level <= 2:
            _add_bottom_border(p)
        return

    if line_type == "list_unordered":
        p = doc.add_paragraph(style="List Bullet")
        _add_inline_formatting(p, data["content"])
        return

    if line_type == "list_ordered":
        p = doc.add_paragraph(style="List Number")
        _add_inline_formatting(p, data["content"])
        return

    if line_type == "blockquote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        _add_inline_formatting(p, data["content"])
        # Apply quote styling to all runs
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(106, 115, 125)
        return

    if line_type == "hr":
        _add_horizontal_rule(doc)
        return

    if line_type == "paragraph":
        p = doc.add_paragraph()
        _add_inline_formatting(p, data["content"])
        return


def _add_inline_formatting(paragraph, text: str):
    """Parse and add inline markdown formatting (bold, italic, code, links, emoji)"""
    # Regex patterns for inline markdown
    patterns = [
        (r"\*\*\*(.+?)\*\*\*", "bold_italic"),
        (r"\*\*(.+?)\*\*", "bold"),
        (r"\*(.+?)\*", "italic"),
        (r"`([^`]+)`", "code"),
        (r"\[([^\]]+)\]\(([^)]+)\)", "link"),
        (r"~~(.+?)~~", "strikethrough"),
    ]

    remaining = text
    segments = []

    while remaining:
        earliest_match = None
        earliest_pos = len(remaining)
        earliest_type = None

        for pattern, fmt_type in patterns:
            match = re.search(pattern, remaining)
            if match and match.start() < earliest_pos:
                earliest_match = match
                earliest_pos = match.start()
                earliest_type = fmt_type

        if earliest_match:
            # Add text before match
            if earliest_pos > 0:
                segments.append(("text", remaining[:earliest_pos]))

            # Add formatted segment
            if earliest_type == "link":
                segments.append(
                    (
                        earliest_type,
                        (earliest_match.group(1), earliest_match.group(2)),
                    )
                )
            else:
                segments.append((earliest_type, earliest_match.group(1)))

            # Continue with remaining text
            remaining = remaining[earliest_match.end() :]
        else:
            # No more matches, add remaining text
            if remaining:
                segments.append(("text", remaining))
            break

    # Render segments
    for seg_type, seg_data in segments:
        if seg_type == "text":
            _add_text_with_emoji_support(paragraph, seg_data)
        elif seg_type == "bold":
            _add_text_with_emoji_support(paragraph, seg_data, bold=True)
        elif seg_type == "italic":
            _add_text_with_emoji_support(paragraph, seg_data, italic=True)
        elif seg_type == "bold_italic":
            _add_text_with_emoji_support(paragraph, seg_data, bold=True, italic=True)
        elif seg_type == "code":
            run = paragraph.add_run(seg_data)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(207, 34, 46)
        elif seg_type == "link":
            link_text, link_url = seg_data
            _add_hyperlink(paragraph, link_url, link_text)
        elif seg_type == "strikethrough":
            _add_text_with_emoji_support(paragraph, seg_data, strikethrough=True)


def _add_text_with_emoji_support(
    paragraph,
    text: str,
    bold: bool = False,
    italic: bool = False,
    strikethrough: bool = False,
):
    """
    Add text to paragraph with proper emoji font support across platforms.

    This function ensures emoji display correctly on:
    - Windows (using Segoe UI Emoji)
    - macOS (using Apple Color Emoji)
    - Linux (using various emoji fonts)

    Strategy: Split text into emoji and non-emoji segments, applying appropriate
    font settings to each segment for cross-platform compatibility.
    """
    if not text:
        return

    # Split text into emoji and non-emoji segments
    segments = []
    current_segment = ""
    is_emoji_segment = False

    for char in text:
        char_is_emoji = char in emoji.EMOJI_DATA

        if char_is_emoji != is_emoji_segment:
            # Segment type changed, save current segment
            if current_segment:
                segments.append((is_emoji_segment, current_segment))
            current_segment = char
            is_emoji_segment = char_is_emoji
        else:
            current_segment += char

    # Add final segment
    if current_segment:
        segments.append((is_emoji_segment, current_segment))

    # Render each segment with appropriate font
    for is_emoji, segment_text in segments:
        run = paragraph.add_run(segment_text)

        # Apply text formatting
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if strikethrough:
            run.font.strike = True

        # Apply emoji-specific font configuration
        if is_emoji:
            _set_emoji_font(run)


def _set_emoji_font(run):
    """
    Configure emoji font with cross-platform fallback support.

    This sets up a font configuration that works across different platforms:
    - WPS Office: Uses Segoe UI Emoji (primary)
    - Windows Word: Uses Segoe UI Emoji (primary)
    - macOS Word: Uses Apple Color Emoji (fallback)
    - Linux: Uses system emoji fonts (fallback)

    Strategy: Explicitly set Segoe UI Emoji for WPS/Windows compatibility,
    while letting macOS/Linux systems fall back to their native emoji fonts.
    """
    try:
        rPr = run._element.get_or_add_rPr()

        # Create font configuration element
        rFonts = OxmlElement("w:rFonts")

        # Explicitly set Segoe UI Emoji for WPS Office and Windows Word
        # This is crucial for WPS which doesn't auto-fallback to emoji fonts
        rFonts.set(qn("w:ascii"), "Segoe UI Emoji")
        rFonts.set(qn("w:hAnsi"), "Segoe UI Emoji")
        rFonts.set(qn("w:eastAsia"), "Segoe UI Emoji")
        rFonts.set(qn("w:cs"), "Segoe UI Emoji")

        # Replace existing font config if present
        existing_rFonts = rPr.find(qn("w:rFonts"))
        if existing_rFonts is not None:
            rPr.remove(existing_rFonts)
        rPr.insert(0, rFonts)

    except Exception as e:
        logger.warning(f"Failed to set emoji font: {e}")


def _add_hyperlink(paragraph, url: str, text: str):
    """Add clickable hyperlink to paragraph"""
    # This is a workaround as python-docx doesn't directly support hyperlinks
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style (blue + underline)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "55B9F7")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def _add_code_block(doc: Document, code: str, language: str):
    """Add code block with background color"""
    p = doc.add_paragraph()

    # Add language label if present
    if language:
        label = p.add_run(f"{language}\n")
        label.font.size = Pt(8)
        label.font.color.rgb = RGBColor(150, 150, 150)

    # Add code content
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_COLOR

    # Add background shading
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), "F6F8FA")
    p._element.get_or_add_pPr().append(shading_elm)

    # Add border
    _add_border(p)


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]):
    """Add markdown table"""
    if not headers:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):
                row_cells[col_idx].text = cell_data


def _add_border(paragraph):
    """Add border around paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), "DCDCDC")
        pBdr.append(border)

    pPr.append(pBdr)


def _add_bottom_border(paragraph):
    """Add bottom border to paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "DCDCDC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_document_footer(doc: Document):
    """Add footer with watermark"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Exported from Wegent")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(160, 160, 160)
